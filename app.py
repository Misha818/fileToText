from flask import Flask, request, jsonify, send_file, render_template
from werkzeug.utils import secure_filename
from docx import Document
from translate import translate_text
from transliterate import transliterate_text
from PIL import Image
import pytesseract
import os
import io
import requests
import base64

# Import your OCR engines here
import easyocr
# from mistral_ocr import run_mistral_ocr

app = Flask(__name__)

# Upload folder
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {"pdf", "png", "jpg", "jpeg"}
MISTRAL_API_KEY=os.getenv('MISTRAL_API_KEY')


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS

def get_extension(filename):
    return filename.rsplit(".", 1)[1].lower()


@app.route("/", methods=["GET"])
def index():
    return render_template("index.html")

@app.route("/upload", methods=["POST"])
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files["file"]
    ocr_engine = request.form.get("engine", "EasyOCR")

    language = request.form.get("language", "en")

    if file.filename == "" or not allowed_file(file.filename):
        return jsonify({"error": "Invalid file type"}), 400
    
    # optimize_image_for_ocr(file)
    filename = secure_filename(file.filename)
    filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    file.save(filepath)

    # OCR processing (placeholders)
    if ocr_engine == "EasyOCR":
        extracted_text = f"[Easy-OCR placeholder] Extracted from {filename}"

        reader = easyocr.Reader([language, 'en'])
        result = reader.readtext(filepath, detail = 0)
        extracted_text = ' '.join(result)

    elif ocr_engine == "MistralOCR":
        api_url = "https://api.mistral.ai/v1/ocr"
        headers = {
            "Authorization": f"Bearer {MISTRAL_API_KEY}",
            "Content-Type": "application/json"
        }
        file_id = upload_file_to_mistral(file, MISTRAL_API_KEY)
        payload = {
            "model": "mistral-ocr-latest",  # Specify the OCR model
            "document": {
                "type": "file",
                "file_id": file_id
            },
            # "language": "hy",
            "pages": [0],  # Process first page, adjust as needed
            "include_image_base64": False,  # Set to True if you need image data back
            "image_limit": 0,
            "image_min_size": 0
        }

        response = requests.post(api_url, headers=headers, json=payload)
        delete_mistral_file(file_id, MISTRAL_API_KEY)

        print(response.status_code)
        
        print(response.text)
        if response.status_code != 200:
            return jsonify({"error": "OCR failed"}), 500

        extracted_text = response.json().get("pages")[0].get("markdown")
    elif ocr_engine == "TesseractOCR":
        image = Image.open(filepath)
        extracted_text = pytesseract.image_to_string(image, lang='ara+eng+rus+hye')

    else:
        extracted_text = f"[Mistral OCR placeholder] Extracted from {filename}"
        # extracted_text = run_mistral_ocr(filepath)

    return jsonify({"text": extracted_text})


@app.route("/translate", methods=["POST"])
def translate():
    data = request.get_json()
    arabic_text = data.get("text", "")
    target_lang = data.get("target_lang", "en")

    if not arabic_text.strip():
        return jsonify({"error": "No text provided"}), 400

    translated = translate_text(arabic_text, target_lang)
    return jsonify({"translation": translated})


@app.route("/transliterate", methods=["POST"])
def transliterate():
    data = request.get_json()
    arabic_text = data.get("text", "")

    if not arabic_text.strip():
        return jsonify({"error": "No text provided"}), 400

    transliterated = transliterate_text(arabic_text)
    return jsonify({"transliteration": transliterated})


@app.route("/download", methods=["POST"])
def download():
    data = request.get_json()
    original = data.get("original", "")
    translation = data.get("translation", "")
    transliteration = data.get("transliteration", "")

    doc = Document()
    doc.add_heading("OCR Output", level=1)

    doc.add_heading("Original Arabic Text", level=2)
    doc.add_paragraph(original)

    if translation:
        doc.add_heading("Translation", level=2)
        doc.add_paragraph(translation)

    if transliteration:
        doc.add_heading("Transliteration", level=2)
        doc.add_paragraph(transliteration)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return send_file(
        buf,
        as_attachment=True,
        download_name="arabic_ocr_output.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

def upload_file_to_mistral(file, api_key):
    """Upload file to Mistral and get file_id"""
    # file = optimize_image_for_ocr(file)
    upload_url = "https://api.mistral.ai/v1/files"
    headers = {
        "Authorization": f"Bearer {api_key}"
    }
    
    file.stream.seek(0)
    files = {
        'file': (file.filename, file.stream, file.content_type),
        'purpose': (None, 'ocr')  # Specify the purpose
    }
    
    response = requests.post(upload_url, headers=headers, files=files)
    
    if response.status_code != 200:
        raise Exception(f"File upload failed: {response.text}")
    
    return response.json()['id']  # Returns the file_id


def delete_mistral_file(file_id, api_key):
    """Delete uploaded file from Mistral"""
    delete_url = f"https://api.mistral.ai/v1/files/{file_id}"
    headers = {
        "Authorization": f"Bearer {api_key}"
    }
    
    response = requests.delete(delete_url, headers=headers)
    return response.status_code == 200


# def optimize_image_for_ocr(file):
#     """Optimize image for better OCR results"""
#     image = Image.open(file.stream)
    
#     # Convert to RGB if needed
#     if image.mode != 'RGB':
#         image = image.convert('RGB')
    
#     # Resize if too small (OCR works better with larger images)
#     width, height = image.size
#     if width < 1000 or height < 1000:
#         scale_factor = max(1000/width, 1000/height)
#         new_size = (int(width * scale_factor), int(height * scale_factor))
#         image = image.resize(new_size, Image.Resampling.LANCZOS)
    
#     # Save optimized image
#     output = io.BytesIO()
#     image.save(output, format='PNG', dpi=(300, 300))
#     output.seek(0)
    
#     return output


if __name__ == "__main__":
    app.run(debug=True)
